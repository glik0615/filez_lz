import docx 
from collections import Counter
import re # импортируем модуль для работы с текстом 
import pandas as pd
import matplotlib.pyplot as plt

a = docx.Document('lion.docx')

b = len(a.paragraphs) # присваиваем переменной значение количества параграфов
word_list = [] #создаем список для неповторяющихся слов

all_letters = []
c = re.compile(r'[а-яё]', re.IGNORECASE) # выражение для поиска русских букв

rus_text = [] #cоздаем список для всего текста (только русские слова)
rus_word = re.compile(r'\b[а-яё]+\b',re.IGNORECASE) #регулярное выражение для поиска русских слов в тексте

for paragraph in a.paragraphs:
    text = paragraph.text.lower() # приводим текст к нижнему регистру
    words = rus_word.findall(text) # находим все русские слова
    letters = c.findall(text) # находим все русские буквы
    all_letters.extend(letters) #добавляем буквы в общий список

    for word in words:
        rus_text.append(word)
        if word not in word_list: # проверка на наличие слова в списке 
            word_list.append(word) #добавляем слово в список

#---------------------------------------------------------------------------------------
#ВСТРЕЧАЕМОСТЬ РУССКИХ СЛОВ В ТЕКСТЕ
rus_text_len = len(rus_text) #присваиваем переменной значение количества слов в списке

word_counts = Counter(rus_text) #подсчитываем частоту встречаемости слов

dict_word = {
    'Слово': list(word_counts.keys()),
    'Частота встречаемости,раз': list(word_counts.values()),
    'Частота встречаемости в %': [round((count / rus_text_len) * 100, 2) for count in word_counts.values()]
}

df_word = pd.DataFrame(dict_word)

doc_word = docx.Document() #создаем новый документ

doc_word.add_heading('Встречаемость слов в тексте', 0)#добавляем заголовок

table = doc_word.add_table(rows=1, cols=3) # Создаем таблицу с одной строкой и тремя колонками 
hdr_cells = table.rows[0].cells#cells используется для доступа к конкретным ячейкам строки, чтобы заполнить их содержимым
hdr_cells[0].text = 'Слово'
hdr_cells[1].text = 'Частота встречи,раз'
hdr_cells[2].text = 'Частота встречи в %'

for i in range(len(df_word)):
    row_cells = table.add_row().cells #добавляем новые строки в таблицу                                         (i name age)
    row_cells[0].text = df_word.loc[i, 'Слово'] #loc выбор строки с индексом [i], то есть в dataframe это по типу (0 bob 12 ) 
    row_cells[1].text = str(df_word.loc[i, 'Частота встречи,раз'])
    row_cells[2].text = str(df_word.loc[i, 'Частота встречи в %'])

doc_word.save('встречаемость_слов.docx') # Сохраняем документ

#--------------------------------------------------------------------------------------
#ВСТРЕЧАЕМОСТЬ БУКВ В ТЕКСТЕ

letters_counts = Counter(all_letters) #подсчитываем частоту встречаемости букв
df_letter = pd.DataFrame(letters_counts.items(),columns=['Буквы','Количество'])
df_letter['Количество'] = pd.to_numeric(df_letter['Количество']) # to_numeric преобразования значений в столбцах DataFrame в числовой тип данных

plt.figure(figsize=(10, 6))# Ширина фигуры будет 10 дюймов, а высота — 6 дюймов.
plt.bar(df_letter['Буквы'], df_letter['Количество'], color='skyblue')#plt.bar создает столбчатую диаграмму.
plt.xlabel('Буквы')#оси X, представляющие буквы
plt.ylabel('Количество')#оси Y, представляющие Количество
plt.title('Встречаемость букв в тексте')#plt.title задает заголовок для графика.
plt.grid(True)#plt.grid(True) включает отображение сетки на графике, что помогает лучше визуализировать данные.
plt.show()
#--------------------------------------------------------------------------------------